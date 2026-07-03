import io
import pdfplumber
import docx
import urllib.request
import json
import mimetypes
import os

try:
    import numpy as np
    from PIL import Image
    import easyocr
    EASYOCR_AVAILABLE = True
except ImportError:
    EASYOCR_AVAILABLE = False

_easyocr_reader = None

def get_ocr_reader():
    global _easyocr_reader
    if not EASYOCR_AVAILABLE:
        return None
    if _easyocr_reader is None:
        # Load reader lazily on first OCR call
        _easyocr_reader = easyocr.Reader(['en'], gpu=False)
    return _easyocr_reader

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extracts text content from a PDF file using pdfplumber."""
    text_content = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)
    except Exception as e:
        print(f"Error extracting PDF: {e}")
    
    return "\n".join(text_content)

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extracts text content from a DOCX file using python-docx."""
    text_content = []
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text:
                text_content.append(para.text)
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Avoid adding identical text if table cells are merged
                    cell_text = cell.text.strip()
                    if cell_text and (not text_content or text_content[-1] != cell_text):
                        text_content.append(cell_text)
    except Exception as e:
        print(f"Error extracting DOCX: {e}")
    
    return "\n".join(text_content)

def extract_text_from_image(file_bytes: bytes) -> str:
    """Extracts text content from an image file using easyocr locally."""
    try:
        image = Image.open(io.BytesIO(file_bytes))
        if image.mode != 'RGB':
            image = image.convert('RGB')
        image_np = np.array(image)
        reader = get_ocr_reader()
        results = reader.readtext(image_np, detail=0)
        return "\n".join(results)
    except Exception as e:
        print(f"Error during local OCR extraction: {e}")
        return ""

def extract_text_from_image_api(file_bytes: bytes, filename: str) -> str:
    """Fallback OCR using OCR.space free API when easyocr/PyTorch is unavailable (e.g. on Render's 512MB RAM free tier)."""
    try:
        api_key = os.environ.get("OCR_SPACE_API_KEY", "helloworld")
        url = 'https://api.ocr.space/parse/image'
        boundary = '----WebKitFormBoundary7MA4YWxkTrZu0gW'
        
        file_basename = os.path.basename(filename)
        mime_type, _ = mimetypes.guess_type(filename)
        if not mime_type:
            mime_type = 'application/octet-stream'
            
        body = []
        body.append(f'--{boundary}'.encode('utf-8'))
        body.append(f'Content-Disposition: form-data; name="apikey"'.encode('utf-8'))
        body.append(''.encode('utf-8'))
        body.append(api_key.encode('utf-8'))
        
        body.append(f'--{boundary}'.encode('utf-8'))
        body.append(f'Content-Disposition: form-data; name="language"'.encode('utf-8'))
        body.append(''.encode('utf-8'))
        body.append('eng'.encode('utf-8'))
        
        body.append(f'--{boundary}'.encode('utf-8'))
        body.append(f'Content-Disposition: form-data; name="file"; filename="{file_basename}"'.encode('utf-8'))
        body.append(f'Content-Type: {mime_type}'.encode('utf-8'))
        body.append(''.encode('utf-8'))
        body.append(file_bytes)
        body.append(f'--{boundary}--'.encode('utf-8'))
        
        request_body = b'\r\n'.join(body)
        headers = {
            'Content-Type': f'multipart/form-data; boundary={boundary}',
            'Content-Length': str(len(request_body))
        }
        
        req = urllib.request.Request(url, data=request_body, headers=headers, method='POST')
        with urllib.request.urlopen(req, timeout=10) as response:
            res_data = response.read().decode('utf-8')
            res_json = json.loads(res_data)
            
            if "ParsedResults" in res_json and len(res_json["ParsedResults"]) > 0:
                return res_json["ParsedResults"][0]["ParsedText"]
            elif "ErrorMessage" in res_json:
                print(f"OCR.space API error: {res_json['ErrorMessage']}")
            else:
                print(f"OCR.space API unexpected response: {res_json}")
    except Exception as e:
        print(f"Failed to perform online OCR fallback: {e}")
    return ""

def parse_resume_bytes(file_bytes: bytes, filename: str) -> str:
    """Detects file type by extension and extracts text."""
    ext = filename.split(".")[-1].lower()
    if ext == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in ["docx", "doc"]:
        return extract_text_from_docx(file_bytes)
    elif ext in ["jpg", "jpeg", "png"]:
        if EASYOCR_AVAILABLE:
            extracted_text = extract_text_from_image(file_bytes)
        else:
            extracted_text = extract_text_from_image_api(file_bytes, filename)
        
        # Validation Heuristics: check if image contains resume keywords
        # Must contain at least two resume sections or keyword markers
        resume_keywords = ["experience", "education", "skills", "projects", "work", "employment", "contact", "summary", "about", "profile"]
        lower_text = extracted_text.lower()
        matched_keywords = [kw for kw in resume_keywords if kw in lower_text]
        
        if len(matched_keywords) < 2 or len(extracted_text.strip()) < 30:
            if not EASYOCR_AVAILABLE and not extracted_text.strip():
                # If online OCR failed or was offline
                raise ValueError("Image OCR engine is currently unavailable. Please upload your resume in PDF or DOCX format.")
            raise ValueError("The uploaded image does not appear to be a valid resume. Please upload a clear image of your resume.")
            
        return extracted_text
    else:
        # Fallback to UTF-8 decoding for plain text files
        try:
            return file_bytes.decode("utf-8")
        except Exception:
            return ""
